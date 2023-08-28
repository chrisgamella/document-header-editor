import tempfile
import unittest
import os
import sys

sys.path.append("..")

from pathlib import Path
import os.path

from document_header_editor import header_generator

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io


folder = os.path.dirname(__file__)
src = folder + os.path.sep + "test_scanned.pdf"
dest = folder + os.path.sep + "test_out_scanned.docx"


def aaa():
    folder = os.path.dirname(__file__)
    a = header_generator.HeaderGenerator()
    print(a.load(folder + os.path.sep + "test.docx"))
    a.set_logo_header(folder + os.path.sep + "logo.png", 1)
    print(a.save(folder + os.path.sep + "test_out.docx"))


def aab():
    
    doc = Document()
    pdf_document = fitz.open(src)
    for page_num in range(pdf_document.page_count):
        pdf_page = pdf_document.load_page(page_num)

        # Create a Pillow image from the bytes
        image = pdf_page.get_pixmap(matrix=fitz.Matrix(300/72, 300/72))
        with tempfile.NamedTemporaryFile(suffix=".png", delete=True, dir=os.getcwd()) as temp_image:
            image.save(temp_image.name, "png")

            # Use pytesseract to extract text from the saved image
            page_text = pytesseract.image_to_string(Image.open(temp_image.name), config='-c preserve_interword_spaces=True')
            
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(page_text)
            
            # Adjust font size and alignment as needed
            font = run.font
            font.size = Pt(12)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(dest)




#app entry
aab()