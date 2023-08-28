from pdf2docx import Converter

from document_header_editor import helper_docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import fitz  # PyMuPDF
import pytesseract
from PIL import Image


def count_characters_in_pdf(pdf_file):
    try:
        # Open the PDF file
        pdf_document = fitz.open(pdf_file)

        # Initialize a variable to count characters
        total_characters = 0

        # Iterate through pages and count characters
        for page_num in range(pdf_document.page_count):
            pdf_page = pdf_document.load_page(page_num)
            page_text = pdf_page.get_text()
            total_characters += len(page_text)

        return total_characters

    except Exception as e:
        print(f"Error: {e}")
        return 0  # Return 0 characters in case of an error
    



def save(src, dest, logo_path, logo_width=1, text_header='', text_font_size=13):
    doc = Document()
    try:
        pdf_document = fitz.open(src)

        for page_num in range(pdf_document.page_count):
            pdf_page = pdf_document.load_page(page_num)

            # Convert the PDF page to an image using PyMuPDF
            image = pdf_page.get_pixmap()
            image_bytes = image.get_bits()

            # Create a Pillow image from the bytes
            pillow_image = Image.frombytes(
                mode="RGB",
                size=(image.width, image.height),
                data=image_bytes,
            )

            # Use pytesseract to extract text from the image
            page_text = pytesseract.image_to_string(pillow_image, config='-c preserve_interword_spaces=True')

            paragraph = doc.add_paragraph()
            run = paragraph.add_run(page_text)
            
            # Adjust font size and alignment as needed
            font = run.font
            font.size = Pt(12)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    except Exception as e:
        pass
        
    # Save the DOCX document
    doc.save(dest)

    return helper_docx.save(dest, dest, logo_path, logo_width=1, text_header='', text_font_size=13)





