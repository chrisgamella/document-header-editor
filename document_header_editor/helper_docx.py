from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.shared import Pt


def save(src, dest, logo_path, logo_width=1, text_header='', text_font_size=13):
    doc = Document(src)
    header = doc.sections[0].header

    # Clear any existing content in the header
    for paragraph in header.paragraphs:
        for run in paragraph.runs:
            run.clear()

    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if logo_path:
        run.add_picture(logo_path, width=Inches(logo_width))  # Adjust width and height as needed
    else:
        run.add_text(text_header)
        run.font.name = "Arial"
        run.font.size = Pt(text_font_size)
        run.bold = True

    # Save the modified DOCX file
    doc.save(dest)

    return dest
